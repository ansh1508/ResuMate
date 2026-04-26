import fitz
import docx
import io


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    filename = filename.lower()
    if filename.endswith(".pdf"):
        return _extract_pdf(content)
    elif filename.endswith(".docx"):
        return _extract_docx(content)
    else:
        raise ValueError("Unsupported file type. Please upload PDF or DOCX.")


def _extract_pdf(content: bytes) -> str:
    text_parts = []
    with fitz.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts).strip()


def _extract_docx(content: bytes) -> str:
    doc = docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs).strip()